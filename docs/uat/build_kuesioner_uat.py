from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Kuesioner-UAT-Ruang-Agunan-Admin-Unit.docx")

GREEN = "005B3A"
GREEN_DARK = "00422C"
GREEN_LIGHT = "EAF4EF"
MINT = "F4F9F6"
GRAY_LIGHT = "F2F4F5"
GRAY_MID = "D8DEE2"
TEXT = "1F2933"
WHITE = "FFFFFF"


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, color=GRAY_MID, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    """Match table width, grid and cells exactly (9360 DXA = 6.5in)."""
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)


def set_font(run, size=11, bold=False, color=TEXT, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, before=0, after=6, line=1.25, alignment=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment


def add_text(paragraph, text, size=11, bold=False, color=TEXT, italic=False):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    return run


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_keep_with_next(p)
    if level == 1:
        add_text(p, text, size=13, bold=True, color=GREEN)
        format_paragraph(p, before=14, after=7, line=1.1)
    else:
        add_text(p, text, size=11.5, bold=True, color=GREEN_DARK)
        format_paragraph(p, before=10, after=5, line=1.1)
    return p


def clear_cell(cell):
    p = cell.paragraphs[0]
    p._element.clear_content()
    return p


def add_label_value_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    set_table_geometry(table, [1700, 7660])
    for row, (label, value) in zip(table.rows, rows):
        for cell in row.cells:
            set_cell_shading(cell, WHITE)
        set_cell_shading(row.cells[0], GREEN_LIGHT)
        p0 = clear_cell(row.cells[0])
        add_text(p0, label, size=10, bold=True, color=GREEN_DARK)
        format_paragraph(p0, after=0, line=1.15)
        p1 = clear_cell(row.cells[1])
        add_text(p1, value, size=10.5, color=TEXT)
        format_paragraph(p1, after=0, line=1.15)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_geometry(table, [9360], indent=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, MINT)
    set_cell_border(cell, color="B9D9C8", size="10")
    set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
    p = clear_cell(cell)
    add_text(p, title + "\n", size=10.5, bold=True, color=GREEN_DARK)
    add_text(p, body, size=10.5, color=TEXT)
    format_paragraph(p, after=0, line=1.25)
    return table


def add_scale_table(doc):
    table = doc.add_table(rows=2, cols=5)
    table.autofit = False
    set_table_geometry(table, [1872] * 5)
    labels = [("5", "Sangat Setuju\n(SS)"), ("4", "Setuju\n(S)"), ("3", "Netral\n(N)"), ("2", "Tidak Setuju\n(TS)"), ("1", "Sangat Tidak Setuju\n(STS)")]
    for index, (score, label) in enumerate(labels):
        cell = table.cell(0, index)
        set_cell_shading(cell, GREEN)
        p = clear_cell(cell)
        add_text(p, score, size=11, bold=True, color=WHITE)
        format_paragraph(p, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        cell = table.cell(1, index)
        set_cell_shading(cell, WHITE)
        p = clear_cell(cell)
        add_text(p, label, size=9.5, color=TEXT)
        format_paragraph(p, after=0, line=1.1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    return table


def add_scenario_table(doc):
    scenarios = [
        ("1", "Akses akun", "Masuk dengan akun Admin Unit dan meninjau menu yang tersedia sesuai peran."),
        ("2", "Kelola barang", "Melihat, menambah atau memperbarui data barang serta media barang menggunakan data uji."),
        ("3", "Pemasaran", "Meninjau kelayakan pemasaran dan membuat pemasaran Harga Tetap atau Lelang Tertutup pada data uji yang sesuai."),
        ("4", "Transaksi", "Meninjau verifikasi bukti pembayaran, status transaksi, dan unggah bukti serah terima pada skenario uji."),
        ("5", "Riwayat & pelanggaran", "Meninjau kronologi barang serta informasi pelanggaran dan pembatasan buyer pada data uji."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    set_table_geometry(table, [600, 2000, 6760])
    headers = ["No.", "Skenario", "Aktivitas yang Dicoba"]
    for idx, label in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, GREEN)
        p = clear_cell(cell)
        add_text(p, label, size=9.5, bold=True, color=WHITE)
        format_paragraph(p, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for no, scenario, activity in scenarios:
        cells = table.add_row().cells
        values = [no, scenario, activity]
        for idx, value in enumerate(values):
            set_cell_shading(cells[idx], WHITE)
            p = clear_cell(cells[idx])
            add_text(p, value, size=9.5, bold=(idx == 1), color=TEXT)
            alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            format_paragraph(p, after=0, line=1.15, alignment=alignment)
    set_table_geometry(table, [600, 2000, 6760])
    return table


def add_question_table(doc, questions):
    widths = [520, 6960, 376, 376, 376, 376, 376]
    table = doc.add_table(rows=1, cols=7)
    table.autofit = False
    headers = ["No.", "Pernyataan", "SS\n5", "S\n4", "N\n3", "TS\n2", "STS\n1"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, GREEN)
        p = clear_cell(cell)
        add_text(p, header, size=8.5 if idx > 1 else 9.5, bold=True, color=WHITE)
        format_paragraph(p, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for number, statement in questions:
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            set_cell_shading(cell, WHITE)
        p = clear_cell(cells[0])
        add_text(p, str(number), size=9.5, bold=True, color=TEXT)
        format_paragraph(p, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        p = clear_cell(cells[1])
        add_text(p, statement, size=9.5, color=TEXT)
        format_paragraph(p, after=0, line=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        for idx in range(2, 7):
            p = clear_cell(cells[idx])
            add_text(p, "□", size=12, color=GREEN_DARK)
            format_paragraph(p, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, widths)
    return table


def add_response_lines(doc, count=4):
    table = doc.add_table(rows=count, cols=1)
    table.autofit = False
    set_table_geometry(table, [9360])
    for row in table.rows:
        cell = row.cells[0]
        set_cell_shading(cell, WHITE)
        set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
        p = clear_cell(cell)
        add_text(p, "", size=10)
        format_paragraph(p, after=0)
    return table


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(paragraph, after=0, line=1.0)
    add_text(paragraph, "Kuesioner UAT Prototipe Sistem Ruang Agunan | Halaman ", size=8.5, color="5E6B73")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    props = doc.core_properties
    props.title = "Kuesioner UAT Prototipe Sistem Ruang Agunan"
    props.subject = "Instrumen User Acceptance Testing untuk responden Admin Unit"
    props.author = "Peneliti"
    props.comments = "Dokumen instrumen UAT untuk kebutuhan penelitian tugas akhir."


def build_document():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    add_text(title, "KUESIONER USER ACCEPTANCE TESTING (UAT)", size=15, bold=True, color=GREEN_DARK)
    format_paragraph(title, before=0, after=2, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    subtitle = doc.add_paragraph()
    add_text(subtitle, "PROTOTIPE SISTEM RUANG AGUNAN", size=12, bold=True, color=GREEN)
    format_paragraph(subtitle, before=0, after=12, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_callout(
        doc,
        "Tujuan Pengisian",
        "Kuesioner ini digunakan untuk menilai penerimaan pengguna terhadap Prototipe Sistem Ruang Agunan dari perspektif Admin Unit. Hasil pengisian hanya digunakan untuk kebutuhan penelitian tugas akhir dan tidak dimaksudkan sebagai persetujuan peluncuran sistem secara resmi.",
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    add_heading(doc, "A. Identitas Responden", level=2)
    add_label_value_table(
        doc,
        [
            ("Kode Responden", "..............................................................."),
            ("Unit Penugasan", "..............................................................."),
            ("Jabatan", "Admin Unit"),
            ("Tanggal Pengisian", "..............................................................."),
        ],
    )
    p = doc.add_paragraph()
    add_text(p, "Catatan: nama responden tidak perlu dicantumkan. Gunakan kode responden untuk menjaga kerahasiaan identitas.", size=9.5, italic=True, color="5E6B73")
    format_paragraph(p, before=4, after=8, line=1.15)

    add_heading(doc, "B. Petunjuk Pengisian", level=2)
    instructions = [
        "Cobalah terlebih dahulu skenario penggunaan yang disiapkan peneliti dengan data uji, bukan data operasional nyata.",
        "Berikan satu jawaban untuk setiap pernyataan dengan memberi tanda centang pada kolom yang paling sesuai.",
        "Jika terdapat kendala atau saran, tuliskan pada bagian komentar di akhir kuesioner.",
    ]
    for number, text in enumerate(instructions, start=1):
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        add_text(p, f"{number}. ", size=10.5, bold=True, color=GREEN_DARK)
        add_text(p, text, size=10.5)
        format_paragraph(p, after=3, line=1.2)

    add_heading(doc, "C. Skala Penilaian", level=2)
    add_scale_table(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    add_heading(doc, "D. Skenario Penggunaan Sebelum Penilaian", level=2)
    p = doc.add_paragraph()
    add_text(p, "Skenario berikut membantu memastikan setiap responden memiliki pengalaman uji yang cukup sebelum memberikan penilaian.", size=10)
    format_paragraph(p, after=4, line=1.15)
    add_scenario_table(doc)

    doc.add_page_break()

    add_heading(doc, "E. Pernyataan Penilaian UAT", level=1)
    p = doc.add_paragraph()
    add_text(p, "Pilih satu respons pada setiap pernyataan berikut.", size=10.5, color="5E6B73")
    format_paragraph(p, after=8, line=1.15)

    functional_questions = [
        (1, "Saya dapat masuk ke halaman Admin Unit dan mengakses fitur yang sesuai dengan hak akses akun saya."),
        (2, "Daftar dan detail barang pada unit penugasan saya ditampilkan dengan jelas sehingga mudah ditelusuri."),
        (3, "Proses menambah atau memperbarui data barang beserta media barang dapat dilakukan dengan mudah menggunakan data uji."),
        (4, "Fitur pengelolaan status barang, seperti perpanjangan jatuh tempo dan penebusan, membantu saya mencatat kondisi barang secara jelas."),
        (5, "Pilihan pemasaran Harga Tetap dan Lelang Tertutup mudah dipahami serta membantu saya menyiapkan pemasaran barang yang memenuhi ketentuan."),
        (6, "Informasi status pemasaran, jadwal lelang, dan nilai transaksi mudah dipantau melalui sistem."),
        (7, "Proses verifikasi atau penolakan bukti pembayaran serta unggah bukti serah terima membantu saya mengelola transaksi secara tertib."),
        (8, "Riwayat barang dan kronologi status membantu saya menelusuri perubahan barang dari waktu ke waktu."),
        (9, "Informasi pelanggaran buyer dan status pembatasannya dapat dilihat serta dipahami dengan jelas oleh Admin Unit."),
    ]
    add_heading(doc, "E.1 Kesesuaian Fungsi", level=2)
    add_question_table(doc, functional_questions)

    doc.add_page_break()

    usability_questions = [
        (10, "Menu dan navigasi pada portal Admin Unit mudah dipahami ketika berpindah dari satu fitur ke fitur lainnya."),
        (11, "Tombol tindakan, formulir, pesan validasi, dan notifikasi membantu saya memahami langkah yang perlu dilakukan."),
        (12, "Informasi pada dashboard, tabel, dan halaman detail tersusun jelas serta mudah dibaca."),
        (13, "Tampilan sistem konsisten sehingga saya tidak mudah bingung saat menggunakan fitur yang berbeda."),
        (14, "Sistem hanya menampilkan data operasional sesuai unit penugasan saya sehingga pengelolaan data terasa lebih aman dan terarah."),
    ]
    acceptance_questions = [
        (15, "Sistem membantu proses pengelolaan barang, pemasaran, dan transaksi menjadi lebih terstruktur dibandingkan pencatatan manual."),
        (16, "Secara keseluruhan, Prototipe Sistem Ruang Agunan dapat diterima untuk mendukung kebutuhan operasional Admin Unit."),
    ]
    add_heading(doc, "E.2 Kemudahan Penggunaan dan Informasi", level=2)
    add_question_table(doc, usability_questions)
    add_heading(doc, "E.3 Manfaat dan Penerimaan Sistem", level=2)
    add_question_table(doc, acceptance_questions)

    add_heading(doc, "F. Komentar dan Saran Responden", level=2)
    p = doc.add_paragraph()
    add_text(p, "Tuliskan masukan terkait fungsi, kemudahan penggunaan, informasi, atau tampilan sistem.", size=10)
    format_paragraph(p, after=4, line=1.15)
    add_response_lines(doc, count=4)

    add_heading(doc, "G. Pengolahan Hasil (Diisi Peneliti)", level=2)
    add_callout(
        doc,
        "Rumus Pengolahan",
        "Skor maksimum = jumlah responden x 16 pernyataan x 5.\nPersentase penerimaan = (total skor yang diperoleh / skor maksimum) x 100%.\nKriteria interpretasi hasil mengikuti pedoman kampus atau rujukan yang digunakan pada bab pengujian.",
    )
    closing = doc.add_paragraph()
    add_text(closing, "Terima kasih atas waktu dan partisipasi Anda.", size=10.5, bold=True, color=GREEN_DARK)
    format_paragraph(closing, before=10, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
